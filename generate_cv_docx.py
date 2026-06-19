from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page Setup ──────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

CONTENT_WIDTH = Inches(7.1)  # 8.5 - 0.7 - 0.7
ACCENT = RGBColor(0x2B, 0x2B, 0x2B)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
BODY = RGBColor(0x33, 0x33, 0x33)
MUTED = RGBColor(0x5F, 0x5F, 0x5F)
LIGHT = RGBColor(0x80, 0x80, 0x80)
RULE_COLOR = "BBBBBB"

# ── Default style ───────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.font.color.rgb = BODY
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.08

rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = doc.element.makeelement(qn("w:rFonts"), {})
    rpr.append(rFonts)
for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
    rFonts.set(qn(attr), "Calibri")


# ── Helpers ─────────────────────────────────────────────────────────────

def _set_bottom_border(paragraph, color=RULE_COLOR, size="6"):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = doc.element.makeelement(qn("w:pBdr"), {})
    bottom = doc.element.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): size,
        qn("w:space"): "1",
        qn("w:color"): color,
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _run(paragraph, text, size=Pt(10), color=BODY, bold=False, italic=False):
    r = paragraph.add_run(text)
    r.font.name = "Calibri"
    r.font.size = size
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    return r


# ── Header Block ────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(1)
_run(p, "MAHER FAYAD", size=Pt(24), color=DARK, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
_run(p, "Senior Product Designer", size=Pt(12), color=MUTED)

# Contact line 1
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(1)
_run(p, "Riyadh, Saudi Arabia", size=Pt(9.5), color=MUTED)
_run(p, "   |   ", size=Pt(9.5), color=LIGHT)
_run(p, "+966 53 584 0297", size=Pt(9.5), color=MUTED)
_run(p, "   |   ", size=Pt(9.5), color=LIGHT)
_run(p, "Contact@maherfayad.com", size=Pt(9.5), color=MUTED)

# Contact line 2
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
_run(p, "maherfayad.com", size=Pt(9.5), color=MUTED)
_run(p, "   |   ", size=Pt(9.5), color=LIGHT)
_run(p, "linkedin.com/in/maherfayad", size=Pt(9.5), color=MUTED)

# Thin rule under header
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
_set_bottom_border(p, color="999999", size="4")


# ── Section Header ──────────────────────────────────────────────────────
def add_section(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(5)
    _run(p, title.upper(), size=Pt(10.5), color=DARK, bold=True)
    # Letter spacing
    r = p.runs[0]._element
    rPr = r.get_or_add_rPr()
    spacing = doc.element.makeelement(qn("w:spacing"), {qn("w:val"): "40"})
    rPr.append(spacing)
    _set_bottom_border(p, color=RULE_COLOR, size="4")


# ── Role Header (company + date on one line, title + location on next) ─
def add_role(company, dates, title, location):
    # Company | Dates
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(CONTENT_WIDTH, alignment=WD_TAB_ALIGNMENT.RIGHT)
    _run(p, company, size=Pt(10.5), color=DARK, bold=True)
    p.add_run("\t")
    _run(p, dates, size=Pt(9.5), color=MUTED)

    # Title | Location
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(3)
    p2.paragraph_format.tab_stops.add_tab_stop(CONTENT_WIDTH, alignment=WD_TAB_ALIGNMENT.RIGHT)
    _run(p2, title, size=Pt(10), color=BODY, italic=True)
    p2.add_run("\t")
    _run(p2, location, size=Pt(9.5), color=LIGHT)


# ── Bullet ──────────────────────────────────────────────────────────────
def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1.5)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    _run(p, "•  ", size=Pt(10), color=LIGHT)
    _run(p, text, size=Pt(10), color=BODY)


# ── Skill Line ──────────────────────────────────────────────────────────
def add_skill(category, skills):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    _run(p, category + ":  ", size=Pt(10), color=DARK, bold=True)
    _run(p, skills, size=Pt(10), color=BODY)


# ── Cert Line ───────────────────────────────────────────────────────────
def add_cert(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1.5)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    _run(p, "•  ", size=Pt(10), color=LIGHT)
    _run(p, text, size=Pt(10), color=BODY)


# ── Education Entry ─────────────────────────────────────────────────────
def add_edu(school, dates, degree, location):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(CONTENT_WIDTH, alignment=WD_TAB_ALIGNMENT.RIGHT)
    _run(p, school, size=Pt(10.5), color=DARK, bold=True)
    p.add_run("\t")
    _run(p, dates, size=Pt(9.5), color=MUTED)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.tab_stops.add_tab_stop(CONTENT_WIDTH, alignment=WD_TAB_ALIGNMENT.RIGHT)
    _run(p2, degree, size=Pt(10), color=BODY, italic=True)
    p2.add_run("\t")
    _run(p2, location, size=Pt(9.5), color=LIGHT)


# ════════════════════════════════════════════════════════════════════════
# CONTENT
# ════════════════════════════════════════════════════════════════════════

# ── PROFESSIONAL SUMMARY ───────────────────────────────────────────────
add_section("Professional Summary")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.18
_run(
    p,
    "Senior Product Designer specializing in bilingual (Arabic/English) experiences "
    "for the GCC market. Designed flows that drove a 47% increase in online account "
    "openings and 81% growth in e-business transactions across Saudi Arabia’s largest "
    "financial institutions. Now designing end-to-end journeys for Almosafer, the GCC’s "
    "leading travel platform, serving millions of users. Deep expertise in design systems, "
    "RTL localization, and analytics-informed funnel optimization.",
    size=Pt(10),
    color=BODY,
)


# ── PROFESSIONAL EXPERIENCE ───────────────────────────────────────────
add_section("Professional Experience")

add_role("Almosafer", "2025 – Present", "Senior Product Designer", "Riyadh, Saudi Arabia")
add_bullet("Design end-to-end Arabic/English booking journeys for Almosafer, the GCC’s leading travel platform serving millions of monthly users across iOS, Android, and web.")
add_bullet("Lead design-system initiatives unifying components and interaction patterns across web and mobile, reducing design-to-dev handoff inconsistency across multiple product surfaces.")
add_bullet("Run analytics-informed funnel optimization (Amplitude) across search, booking, and post-trip flows, identifying drop-off points and shipping design fixes to improve conversion.")
add_bullet("Localize and design RTL-first Arabic experiences using GCC conventions, partnering with PMs and engineers to ship research-validated flows.")

add_role("AZMX (Al Rajhi Bank)", "Dec 2023 – 2025", "Sr. User Experience Designer", "Riyadh, Saudi Arabia")
add_bullet("Delivered 41 revamps and enhancements to Al Rajhi e-business platforms in Figma, improving usability for enterprise and retail banking customers.")
add_bullet("Boosted online account openings by 47% through redesigned onboarding user flows and reduced friction points.")
add_bullet("Increased transactions via e-business channels by 81% by streamlining key task flows, navigation, and design-system consistency.")
add_bullet("Redesigned the enterprise payroll system with bilingual RTL Arabic/English interfaces and analytics-informed user journeys.")
add_bullet("Designed across multiple Al Rajhi group entities including Al Rajhi Capital, Emkan Finance, and neoleap.")

add_role("Contact Financial Holding", "Dec 2022 – Nov 2023", "Product Designer", "Cairo, Egypt")
add_bullet("Owned the full design lifecycle from research to Figma handoff for 2 core fintech products: Contact Now and Contact Brokerage.")
add_bullet("Conducted user research, interviews, and usability testing across 3 product verticals to validate design decisions and reduce development rework.")
add_bullet("Built and maintained a component library in Figma to standardize UI patterns across the product suite.")
add_bullet("Shipped research-driven user flows and prototypes that aligned product, engineering, and business stakeholders on scope.")

add_role(
    "Freelance Product Designer (via Upwork)",
    "Aug 2022 – Dec 2023",
    "Clients: SuperSight, Konica Minolta, MIT Olin Foundation, IterationX, Theradome, Solidity Law, Six Clovers",
    "Remote",
)
add_bullet("Designed Sanarte, a mobile soundscape app for remote workers, achieving 100% task completion rate in usability testing across 5 core user flows.")
add_bullet("Designed the LFG habit-tracking app for the MIT Olin Foundation using Octalysis gamification, increasing user engagement by 20%.")
add_bullet("Delivered end-to-end product design (research, wireframes, Figma prototypes, usability testing) for 15+ clients across US, Europe, and MENA.")
add_bullet("Designed Campus51, an LMS platform with end-to-end educator journeys for curriculum creation, student engagement, and professional training.")
add_bullet("Designed Deployo, an AI model deployment dashboard in Figma for DevOps teams, unifying integration pipelines and performance monitoring.")
add_bullet("Created Six Clovers, a decentralized payment portal with developer documentation and corporate finance transaction flows.")

add_role("GameIT", "Jun 2022 – Dec 2022", "UI/UX Designer (Part-time)", "Remote")
add_bullet("Redesigned website user flows for usability and engagement, improving customer satisfaction through iterative design and testing.")
add_bullet("Partnered with the game development team to optimize accessibility and interaction patterns across game interfaces.")
add_bullet("Created visual and video assets with the marketing team, strengthening brand presence across digital channels.")

add_role("Algoriza", "Mar 2022 – Jun 2022", "UX Design Intern", "Remote")
add_bullet("Completed end-to-end UX training: user interviews, persona development, journey mapping, and target-customer analysis.")
add_bullet("Contributed wireframes and design documentation under senior designer mentorship.")


# ── CERTIFICATIONS ─────────────────────────────────────────────────────
add_section("Certifications")
for c in [
    "Google UX Design Professional Certificate  (Google / Coursera)",
    "Google Data Analytics Professional Certificate  (Google / Coursera)",
    "Enterprise Design Thinking Practitioner  (IBM)",
    "Enterprise Design Thinking Co-Creator  (IBM)",
    "McKinsey Forward Program  (McKinsey & Company)",
    "Meta Front-End Developer Certificate  (Meta / Coursera)",
    "Product Analytics Certification  (Pendo / Mind the Product)",
]:
    add_cert(c)


# ── SKILLS ─────────────────────────────────────────────────────────────
add_section("Skills")
add_skill("Core", "Product Design, Design Systems, Interaction Design, Information Architecture, Bilingual & RTL Design (Arabic/English), Accessibility (WCAG 2.2)")
add_skill("Research & Analytics", "User Research, Usability Testing, Journey Mapping, A/B Testing, Amplitude, Product Analytics, Hotjar")
add_skill("Tools", "Figma, FigJam, Sketch, Adobe XD, Axure, Miro")
add_skill("Methods", "Design Thinking, Prototyping, Agile/Scrum, Octalysis Gamification")
add_skill("Development", "HTML, CSS, React (fundamentals)")
add_skill("Languages", "Arabic (Native), English (Professional Working Proficiency)")


# ── EDUCATION ──────────────────────────────────────────────────────────
add_section("Education")
add_edu("Mansoura University", "2014 – 2019", "Bachelor of Science in Mechatronics Engineering", "Mansoura, Egypt")


# ════════════════════════════════════════════════════════════════════════
# Save
# ════════════════════════════════════════════════════════════════════════
out = r"c:\Users\Admin\Documents\GitHub\New folder\Maher-Fayad-Senior-Product-Designer.docx"
doc.save(out)
print(f"Done  ->  {out}")
